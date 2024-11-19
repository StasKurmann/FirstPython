from tkinter import *  # Erzeugung des Fensters
import random  # Improt Method random für Zufahlszahlengenerator
from docx import Document


def zufallszahl(laenge):  # Definierung Funktion Zufallzahl
    # Gibt die Zahl von 1 bis Länge zurück
    return random.randrange(1, laenge)


def Word_Generator(Tabelle, Dateiname):  # Eine Funktion zum Generieren von Word-Datei

    document = Document()

    document.add_heading('Aufgaben zum Rechnen', 0)

    document.add_paragraph(
        'Schreibe richtige Antworten in der Spalte "Ergebnis" ')

    records = Tabelle

    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Zahl'
    hdr_cells[1].text = 'Operator'
    hdr_cells[2].text = 'Zahl'
    hdr_cells[3].text = 'Ergebnis'

    for Zahl, Operator, Zahl2, Ergebnis in records:
        row_cells = table.add_row().cells
        row_cells[0].text = str(Zahl)
        row_cells[1].text = Operator
        row_cells[2].text = str(Zahl2)
        row_cells[3].text = str(Ergebnis)

    # document.add_page_break()

    document.save(Dateiname)


# Definierung Funktion Addieren, Subtrahieren, Multiplizieren


def Funktion_add_sub_mult(_operator_funktion_, _laenge_add_sub_mult_, _Anzahl_Rechnenbeispiele_):
    """Diese Funktion generiert beliebige Anzahl von Rechenbeispielen und schreibt in Form einer 
    Tabelle die Beispiele in einer txt.Datei. 
    Eingangswerte sind:
    Operator = als String
    Anzahl Beispiele = als Int
    Zahlenbereich in dem die Zahlen generiert werden.
    """
    A = []  # leerere Liste
    x = 0
    # for-Schleife
    for x in range(_Anzahl_Rechnenbeispiele_):
        B = [(zufallszahl(_laenge_add_sub_mult_), _operator_funktion_, zufallszahl(
            _laenge_add_sub_mult_), "= ____________")]  # temporäre Liste
        A.extend(B)
        x += 1
    return A


# Definierung Funktion Dividieren
def Funktion_div(_laenge_div_, _Anzahl_Rechnenbeispiele_):
    """Diese Funktion generiert beliebige Anzahl von Rechenbeispielen mit Dividieren und schreibt in Form einer 
    Tabelle die Beispiele in einer txt.Datei. 
    Eingangswerte sind:
    Anzahl Beispiele = als Int
    Zahlenbereich in dem die Zahlen generiert werden
    Dateiname als Name der Datei, in der die Tabelle gespeichert werden muss.
    """

    x = 0  # Provisorische Zahl x für While-Schleife
    A = []
    # while-Schleife
    while x < _Anzahl_Rechnenbeispiele_:
        divident = zufallszahl(_laenge_div_)
        divisor = zufallszahl(_laenge_div_)
        if divident % divisor == 0 and divisor != 1:  # Prüft, ob Divident durch divisor teilbar ohne Rest ist
            # weil die Aufgaben im Kopf rechenbar sein müssen. Weiterhin wird auch geprüft, ob Divisor ungleich 1 ist, da durch 1 einfach zu teilen ist.
            B = [(divident, "/", divisor, "= ____________")]
            A.extend(B)
            x += 1
    return A


# Eine Funktion definieren, wo alle Generatorfunktionen hinterlegt sind.
def Generatorfunktionen():
 # Aufruf aller Funktionen
    # Operator, Zahlenbereich, Anzahl von Rechenaufgaben,
    # Name der Datei, wo die Aufgaben geschrieben werden
    Word_Generator(Funktion_add_sub_mult("+", int(entryZahl_bis_Addiere.get()),
                                         int(entryAnzahl_Aufgaben_Addiere.get()),), "Addieren.docx")
    Word_Generator(Funktion_add_sub_mult("-", int(entryZahl_bis_Subtrahieren.get()),
                                         int(entryAnzahl_Aufgaben_Subtrahieren.get())), "Subtrahieren.docx")
    Word_Generator(Funktion_add_sub_mult("*", int(entryZahl_bis_Multiplizieren.get()),
                                         int(entryAnzahl_Aufgaben_Multiplizieren.get())), "Multiplizieren.docx")
    Word_Generator(Funktion_div(int(entryZahl_bis_Dividieren.get()),
                                int(entryAnzahl_Aufgaben_Dividieren.get())), "Dividieren.docx")


def buttonGenerierenClick():  # Definieren vom Button "Generieren".
    Generatorfunktionen()


# Initialisierung TKinter
tkFenster = Tk()
tkFenster.title("Rechenbeispiele")
tkFenster.geometry('500x300')

# Label für die Anzeige Tittel
labelTitle = Label(master=tkFenster, text="Rechenaufgaben",
                   fg="white", bg="gray", font=("Aptos", 16))
labelTitle.place(x=150, y=5, width=200, height=20)

# Label für die Anzeige der Daten für Addieren
labelUntertitle_Addieren = Label(
    master=tkFenster, text="Für Addieren", fg="black", bg="white", font=("Aptos", 11))
labelUntertitle_Addieren.place(x=10, y=30, width=190, height=20)

# Label von 1 bis ... für Addieren
label_von_bis_Addieren = Label(
    master=tkFenster, text="von 1 bis", font=("Aptos", 10))
label_von_bis_Addieren.place(x=10, y=60, width=50, height=20)

# Entry für von - bis Addieren
entryZahl_bis_Addiere = Entry(master=tkFenster, bg="white")
entryZahl_bis_Addiere.place(x=120, y=60, width=40, height=20)

# Label Anzahl Aufgaben für Addieren
label_Anzahl_Aufgaben_Addieren = Label(
    master=tkFenster, text="Anzahl Aufgaben", font=("Aptos", 10))
label_Anzahl_Aufgaben_Addieren.place(x=10, y=90, width=100, height=20)

# Entry für Anzahl Aufgaben Addieren
entryAnzahl_Aufgaben_Addiere = Entry(master=tkFenster, bg="white")
entryAnzahl_Aufgaben_Addiere.place(x=120, y=90, width=40, height=20)

# Label für die Anzeige der Daten für Subtrahieren
labelUntertitle_Subtrahieren = Label(
    master=tkFenster, text="Für Subtrahieren", fg="black", bg="white", font=("Aptos", 11))
labelUntertitle_Subtrahieren.place(x=300, y=30, width=190, height=20)

# Label von 1 bis ... für Subtrahieren
label_von_bis_Subtrahieren = Label(
    master=tkFenster, text="von 1 bis", font=("Aptos", 10))
label_von_bis_Subtrahieren.place(x=300, y=60, width=50, height=20)

# Entry für von - bis Subtrahieren
entryZahl_bis_Subtrahieren = Entry(master=tkFenster, bg="white")
entryZahl_bis_Subtrahieren.place(x=410, y=60, width=40, height=20)

# Label Anzahl Aufgaben für Subtrahieren
label_Anzahl_Aufgaben_Subtrahieren = Label(
    master=tkFenster, text="Anzahl Aufgaben", font=("Aptos", 10))
label_Anzahl_Aufgaben_Subtrahieren.place(x=300, y=90, width=100, height=20)

# Entry für Anzahl Aufgaben Subtrahieren
entryAnzahl_Aufgaben_Subtrahieren = Entry(master=tkFenster, bg="white")
entryAnzahl_Aufgaben_Subtrahieren.place(x=410, y=90, width=40, height=20)


# Label für die Anzeige der Daten für Multiplizieren
labelUntertitle_Multiplizieren = Label(
    master=tkFenster, text="Für Multiplizieren", fg="black", bg="white", font=("Aptos", 11))
labelUntertitle_Multiplizieren.place(x=10, y=120, width=190, height=20)

# Label von 1 bis ... für Multiplizieren
label_von_bis_Multiplizieren = Label(
    master=tkFenster, text="von 1 bis", font=("Aptos", 10))
label_von_bis_Multiplizieren.place(x=10, y=150, width=50, height=20)

# Entry für von - bis Multiplizieren
entryZahl_bis_Multiplizieren = Entry(master=tkFenster, bg="white")
entryZahl_bis_Multiplizieren.place(x=120, y=150, width=40, height=20)

# Label Anzahl Aufgaben für Multiplizieren
label_Anzahl_Aufgaben_Multiplizieren = Label(
    master=tkFenster, text="Anzahl Aufgaben", font=("Aptos", 10))
label_Anzahl_Aufgaben_Multiplizieren.place(x=10, y=180, width=100, height=20)

# Entry für Anzahl Aufgaben Multiplizieren
entryAnzahl_Aufgaben_Multiplizieren = Entry(master=tkFenster, bg="white")
entryAnzahl_Aufgaben_Multiplizieren.place(x=120, y=180, width=40, height=20)

# Label für die Anzeige der Daten für Dividieren
labelUntertitle_Dividieren = Label(
    master=tkFenster, text="Für Dividieren", fg="black", bg="white", font=("Aptos", 11))
labelUntertitle_Dividieren.place(x=300, y=120, width=190, height=20)

# Label von 1 bis ... für Dividieren
label_von_bis_Dividieren = Label(
    master=tkFenster, text="von 1 bis", font=("Aptos", 10))
label_von_bis_Dividieren.place(x=300, y=150, width=50, height=20)

# Entry für von - bis Dividieren
entryZahl_bis_Dividieren = Entry(master=tkFenster, bg="white")
entryZahl_bis_Dividieren.place(x=410, y=150, width=40, height=20)

# Label Anzahl Aufgaben für Dividieren
label_Anzahl_Aufgaben_Dividieren = Label(
    master=tkFenster, text="Anzahl Aufgaben", font=("Aptos", 10))
label_Anzahl_Aufgaben_Dividieren.place(x=300, y=180, width=100, height=20)

# Entry für Anzahl Aufgaben Diviedieren
entryAnzahl_Aufgaben_Dividieren = Entry(master=tkFenster, bg="white")
entryAnzahl_Aufgaben_Dividieren.place(x=410, y=180, width=40, height=20)

# Button zum Generieren
buttonGenerieren = Button(master=tkFenster, bg='#FBD975',
                          text="Generieren", font=("Aptos", 12), command=buttonGenerierenClick)
buttonGenerieren.place(x=200, y=220, width=100, height=27)

# Button zum Schließen
buttonQuit = Button(master=tkFenster, bg="gray",
                    text="Quit", font=("Aptos", 12), command=tkFenster.destroy)
buttonQuit.place(x=400, y=260, width=50, height=27)


tkFenster.mainloop()  # Aktivierung des Fensters
